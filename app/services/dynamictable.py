from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, List, Any, Optional

MAX_DYNAMIC_COLS = 16
STAFF_FONT_SIZE = Pt(12)
DEFAULT_FONT_SIZE = Pt(10)
HEADER_FONT_SIZE = Pt(9.5)

SERIAL_COL_WIDTH = Cm(0.9)
STAFF_COL_WIDTH = Cm(5.0)
MONTH_COL_WIDTH = Cm(1.3)
TOTAL_COL_WIDTH = Cm(1.6)    


def set_cell_shading(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag == qn("w:shd"):
            tcPr.remove(child)

    normalized = str(color_hex).strip().strip("'").strip('"').replace("#", "")
    if not normalized:
        return

    shd = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), normalized))
    tcPr.append(shd)

def set_cell_text_bold(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

def set_cell_font(cell, size: Pt, bold=False, name="Calibri"):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = size
            r.font.bold = bold
            r.font.name = name

def align_cell(cell, align='center'):
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT
    }
    for p in cell.paragraphs:
        p.alignment = alignment_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

def set_cell_vertical_alignment(cell, valign='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), valign)
    tcPr.append(vAlign)

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcMar = OxmlElement('w:tcMar')
    
    if top > 0:
        top_mar = OxmlElement('w:top')
        top_mar.set(qn('w:w'), str(int(top * 20)))
        top_mar.set(qn('w:type'), 'dxa')
        tcMar.append(top_mar)
    
    if bottom > 0:
        bottom_mar = OxmlElement('w:bottom')
        bottom_mar.set(qn('w:w'), str(int(bottom * 20))) 
        bottom_mar.set(qn('w:type'), 'dxa')
        tcMar.append(bottom_mar)
    
    if left > 0:
        left_mar = OxmlElement('w:left')
        left_mar.set(qn('w:w'), str(int(left * 20)))  
        left_mar.set(qn('w:type'), 'dxa')
        tcMar.append(left_mar)
    
    if right > 0:
        right_mar = OxmlElement('w:right')
        right_mar.set(qn('w:w'), str(int(right * 20)))  
        right_mar.set(qn('w:type'), 'dxa')
        tcMar.append(right_mar)
    
    if tcMar.getchildren(): 
        tcPr.append(tcMar)

def set_cell_width(cell, width_cm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm.cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_row_height(row, height_cm):
    row_elem = row._tr
    trPr = row_elem.get_or_add_trPr()

    for child in list(trPr):
        if child.tag == qn("w:trHeight"):
            trPr.remove(child)

    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm.cm * 567)))
    trHeight.set(qn('w:type'), 'dxa')
    trPr.append(trHeight)



def set_table_width_and_alignment(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    tblJc = OxmlElement('w:jc')
    tblJc.set(qn('w:val'), 'center')
    tblPr.append(tblJc)
    
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

def enforce_column_widths(table, num_static_cols=2, num_month_cols=0, has_total=False):
    try:
        if len(table.columns) > 1:
            table.columns[0].width = SERIAL_COL_WIDTH
            table.columns[1].width = STAFF_COL_WIDTH

            for i in range(2, len(table.columns)):
                if has_total and i == len(table.columns) - 1:
                    table.columns[i].width = TOTAL_COL_WIDTH
                else:
                    table.columns[i].width = MONTH_COL_WIDTH
    except:
        pass
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if col_idx == 0:
                set_cell_width(cell, SERIAL_COL_WIDTH)
            elif col_idx == 1:
                set_cell_width(cell, STAFF_COL_WIDTH)
            elif has_total and col_idx == len(row.cells) - 1:
                set_cell_width(cell, TOTAL_COL_WIDTH)
            else:
                set_cell_width(cell, MONTH_COL_WIDTH)

def add_row_with_color(table, values, color=None):
    row = table.add_row().cells
    for i, v in enumerate(values):
        row[i].text = str(v) if v is not None else ""
        if color:
            set_cell_shading(row[i], color)

def _generate_legend_table(doc, legend, color_map):
    if not legend:
        return None
    
    legend_table = doc.add_table(rows=0, cols=1)
    legend_table.style = "Table Grid"
    set_table_width_and_alignment(legend_table)
    
    try:
        legend_table.columns[0].width = Cm(4.0)
    except:
        pass

    for legend_item in legend:
        phase_name = legend_item.get('phase', '')
        phase_color = legend_item.get('color', color_map.get(phase_name.strip().lower(), '#FFFFFF'))
        
        row = legend_table.add_row()
        cell = row.cells[0]
        
        cell.text = phase_name
        if phase_color:
            set_cell_shading(cell, phase_color)
        set_cell_font(cell, DEFAULT_FONT_SIZE, True)
        align_cell(cell, 'center')
    
    return legend_table

def find_sdt_by_title(doc: Document, title: str):
    for el in doc.element.body.iter():
        if el.tag.endswith("sdt"):
            for child in el.iter():
                if child.tag.endswith("alias") and child.get(qn("w:val")) == title:
                    return el
                if child.tag.endswith("tag") and child.get(qn("w:val")) == title:
                    return el
    return None


def parse_color_mapping(colors):
    cmap = {}
    for c in colors or []:
        for k, v in c.items():
            key = k.strip().lower()
            if isinstance(v, dict):
                cmap[key] = v.get("color")
            else:
                cmap[key] = v
    return cmap
def detect_table_format(rows, headers, colors):
    if not rows:
        return 'simple'
    if 'phaseType' in rows[0] or 'phase' in rows[0]:
        return 'phase_based'
    if headers:
        return 'custom'
    return 'simple'


def generate_dynamic_table(doc: Document, tag: str, data: Dict[str, Any]) -> Document:
    headers = data.get("headers", [])
    rows = data.get("rows", [])
    colors = data.get("colors", [])
    legend = data.get("legend", [])
    if "headerColor" in data and data.get("headerColor") is not None:
        header_color = data.get("headerColor")
    else:
        header_color = "#333399"

    if not rows:
        return doc

    color_map = parse_color_mapping(colors)
    table_format = detect_table_format(rows, headers, colors)

    sdt = find_sdt_by_title(doc, tag)
    if sdt is None:
        return doc

    sdt_content = next(el for el in sdt.iter() if el.tag.endswith("sdtContent"))
    
    for child in list(sdt_content):
        sdt_content.remove(child)

    if table_format == 'phase_based':
        table = _generate_phase_based_table(doc, rows, color_map, header_color)
        sdt_content.append(table._tbl)

    elif table_format == 'custom':
        max_cols = data.get("customMaxCols", MAX_DYNAMIC_COLS)
        tables = _generate_custom_table(doc, headers, rows, color_map, header_color, legend, max_cols)

        for i, table in enumerate(tables):
            if i > 0:
                spacing_p = doc.add_paragraph()
                spacing_p.space_before = Pt(12)
                spacing_p.space_after = Pt(6)
                sdt_content.append(spacing_p._element)
            sdt_content.append(table._tbl)

    else:
        table = _generate_simple_table(doc, rows, color_map, header_color)
        sdt_content.append(table._tbl)

    if legend:
        legend_spacing_p = doc.add_paragraph()
        legend_spacing_p.space_before = Pt(18)
        legend_spacing_p.space_after = Pt(6)
        sdt_content.append(legend_spacing_p._element)
        
        legend_table = _generate_legend_table(doc, legend, color_map)
        if legend_table:
            sdt_content.append(legend_table._tbl)

    return doc


def _generate_phase_based_table(doc, rows, color_map, header_color):
    headers = [k for k in rows[0] if k not in ('phaseType', 'phase')]
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"
    
    set_table_width_and_alignment(table)
    
    header_row = table.add_row()
    header_cells = header_row.cells

    set_row_height(header_row, Cm(0.6))
    
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = str(header)
        set_cell_shading(cell, header_color)
        set_cell_text_bold(cell)
        set_cell_font(cell, HEADER_FONT_SIZE, True)

        if idx == 0:
            align_cell(cell, 'center')
        elif idx == 1:
            align_cell(cell, 'left')
        else:
            align_cell(cell, 'center')
    
    for r in rows:
        data_row = table.add_row()
        data_cells = data_row.cells
        phase = r.get('phaseType') or r.get('phase')
        phase_key = phase.strip().lower() if isinstance(phase, str) else None
        
        for idx, h in enumerate(headers):
            cell = data_cells[idx]
            cell.text = str(r.get(h, ''))

            if idx == 0:
                set_cell_font(cell, DEFAULT_FONT_SIZE, False)
                align_cell(cell, 'center')
            elif idx == 1:
                set_cell_font(cell, STAFF_FONT_SIZE, False)
                align_cell(cell, 'left')
            else:
                set_cell_font(cell, DEFAULT_FONT_SIZE, False)
                align_cell(cell, 'center')
            
            if phase_key and phase_key in color_map:
                set_cell_shading(cell, color_map[phase_key])
    
    enforce_column_widths(table, 2, len(headers) - 2, False)
    
    return table


def _generate_simple_table(doc, rows, color_map, header_color):
    headers = list(rows[0].keys())
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = "Table Grid"
    
    set_table_width_and_alignment(table)
    
    header_row = table.add_row()
    header_cells = header_row.cells

    set_row_height(header_row, Cm(0.6))
    
    for idx, header in enumerate(headers):
        cell = header_cells[idx]
        cell.text = str(header)
        set_cell_shading(cell, header_color)
        set_cell_text_bold(cell)
        set_cell_font(cell, HEADER_FONT_SIZE, True)

        if idx == 0:
            align_cell(cell, 'center')
        elif idx == 1:
            align_cell(cell, 'left')
        else:
            align_cell(cell, 'center')
    
    for r in rows:
        data_row = table.add_row()
        data_cells = data_row.cells
        
        for idx, h in enumerate(headers):
            cell = data_cells[idx]
            cell.text = str(r.get(h, ''))

            if idx == 0:
                set_cell_font(cell, DEFAULT_FONT_SIZE, False)
                align_cell(cell, 'center')
            elif idx == 1:
                set_cell_font(cell, STAFF_FONT_SIZE, False)
                align_cell(cell, 'left')
            else:
                set_cell_font(cell, DEFAULT_FONT_SIZE, False)
                align_cell(cell, 'center')
    
    enforce_column_widths(table, 2, len(headers) - 2, False)
    
    return table


def _generate_custom_table(doc, headers, rows, color_map, header_color, legend, max_dynamic_cols=MAX_DYNAMIC_COLS):
    all_months_set = set()
    for row in rows:
        for month_obj in row.get('months', []):
            all_months_set.update(month_obj.keys())

    month_headers = [h for h in headers if h in all_months_set]
    static_headers = [h for h in headers if h not in all_months_set]
    has_total = 'Total' in static_headers
    tables = []
    
    if len(month_headers) > max_dynamic_cols:
        is_paginated_table = True
        for chunk_idx in range(0, len(month_headers), max_dynamic_cols):
            chunk = month_headers[chunk_idx:chunk_idx + max_dynamic_cols]
            row_offset = chunk_idx + 1
            remaining_months_after_this_chunk = len(month_headers) - (chunk_idx + len(chunk))
            is_last_chunk = remaining_months_after_this_chunk == 0
            
            if has_total and is_last_chunk and len(chunk) < max_dynamic_cols:
                table_headers = static_headers[:2] + chunk + ['Total']
                table = _create_single_table(doc, table_headers, rows, color_map, header_color, 
                                           chunk, True, is_paginated_table, row_offset)
            else:
                table_headers = static_headers[:2] + chunk
                table = _create_single_table(doc, table_headers, rows, color_map, header_color, 
                                           chunk, False, is_paginated_table, row_offset)
            
            tables.append(table)
        
        if has_total and len(month_headers) % max_dynamic_cols == 0:
            total_table_headers = static_headers[:2] + ['Total']
            total_table = _create_single_table(doc, total_table_headers, rows, color_map, header_color, 
                                             [], True, True, 0)
            tables.append(total_table)
    else:
        table = _create_single_table(doc, headers, rows, color_map, header_color, 
                                   month_headers, has_total, False, 1)
        tables.append(table)
    
    return tables

def _create_single_table(doc, table_headers, rows, color_map, header_color, 
                        month_chunk, has_total_in_this_table, is_split_table, row_offset=0):
    table = doc.add_table(rows=0, cols=len(table_headers))
    table.style = "Table Grid"
    set_table_width_and_alignment(table)
    total_col_idx = len(table_headers) - 1 if has_total_in_this_table and table_headers else None

    numbering_row = table.add_row()
    numbering_cells = numbering_row.cells
    set_row_height(numbering_row, Cm(0.5))
    
    month_col_counter = row_offset
    for idx, header_text in enumerate(table_headers):
        cell = numbering_cells[idx]
        header_key = str(header_text).strip().lower()
        is_total_col = header_key == 'total' or (total_col_idx is not None and idx == total_col_idx)
        
        if idx < 2:
            cell.text = ""
            set_cell_shading(cell, header_color)
        elif is_total_col:
            cell.text = ""
            set_cell_shading(cell, '#FFF2CC')
        else:
            cell.text = str(month_col_counter)
            month_col_counter += 1
            set_cell_shading(cell, header_color)
        
        set_cell_font(cell, Pt(8), True)
        align_cell(cell, 'center')
        set_cell_vertical_alignment(cell, 'center')

    header_row = table.add_row()
    header_cells = header_row.cells
    
    # Set header row height to be reduced
    set_row_height(header_row, Cm(0.6))

    for idx, header_text in enumerate(table_headers):
        cell = header_cells[idx]
        header_key = str(header_text).strip().lower()
        is_total_col = header_key == 'total' or (total_col_idx is not None and idx == total_col_idx)

        if idx < 2 or is_total_col:
            merged_cell = numbering_cells[idx].merge(cell)
            merged_cell.text = str(header_text)
            if is_total_col:
                set_cell_shading(merged_cell, '#FFF2CC')
            else:
                set_cell_shading(merged_cell, header_color)
            set_cell_text_bold(merged_cell)
            set_cell_font(merged_cell, HEADER_FONT_SIZE, True)
            set_cell_vertical_alignment(merged_cell, 'center')
            if idx == 0:
                align_cell(merged_cell, 'center')
            elif idx == 1:
                align_cell(merged_cell, 'left')
            else:
                align_cell(merged_cell, 'center')
            continue

        cell.text = str(header_text)
        if is_total_col:
            set_cell_shading(cell, '#FFF2CC')
        else:
            set_cell_shading(cell, header_color)
        set_cell_text_bold(cell)
        set_cell_font(cell, HEADER_FONT_SIZE, True)
        set_cell_vertical_alignment(cell, 'center')

        if idx == 0:
            align_cell(cell, 'center')
        elif idx == 1:
            align_cell(cell, 'left')
        else:
            align_cell(cell, 'center')

    for row_idx, row_data in enumerate(rows):
        data_row = table.add_row()
        set_row_height(data_row, Cm(0.5))
        data_cells = data_row.cells

        month_map = {}
        for month_obj in row_data.get('months', []):
            for month_name, month_info in month_obj.items():
                month_map[month_name] = month_info

        for col_idx, header in enumerate(table_headers):
            cell = data_cells[col_idx]
            header_key = str(header).strip().lower()
            is_total_col = header_key == 'total' or (total_col_idx is not None and col_idx == total_col_idx)

            if col_idx == 0:
                cell_value = str(row_idx + 1)
            elif col_idx == 1:
                cell_value = row_data.get(header, '')
            elif is_total_col:
                cell_value = row_data.get(header, '')
            else:
                month_info = month_map.get(header, {})
                month_value = month_info.get('value', '')
                cell_value = month_value if str(month_value).strip() else '-'

            cell.text = str(cell_value)

            if col_idx == 0:
                set_cell_font(cell, DEFAULT_FONT_SIZE, False)
                align_cell(cell, 'center')
                set_cell_vertical_alignment(cell, 'center')
            elif col_idx == 1:
                set_cell_font(cell, STAFF_FONT_SIZE, False)
                align_cell(cell, 'left')
                set_cell_vertical_alignment(cell, 'top')
            elif is_total_col:
                set_cell_font(cell, DEFAULT_FONT_SIZE, True)
                align_cell(cell, 'center')
                set_cell_vertical_alignment(cell, 'top')
                set_cell_shading(cell, '#FFF2CC')
            else:
                phase = month_info.get('phase')
                phase_key = phase.strip().lower() if isinstance(phase, str) else None
                if phase_key and phase_key in color_map:
                    set_cell_shading(cell, color_map[phase_key])
                set_cell_font(cell, DEFAULT_FONT_SIZE, False)
                align_cell(cell, 'center')
                set_cell_vertical_alignment(cell, 'top')

    month_maps = []
    for row_data in rows:
        month_map = {}
        for month_obj in row_data.get('months', []):
            for month_name, month_info in month_obj.items():
                month_map[month_name] = month_info
        month_maps.append(month_map)

    def _to_float(value):
        try:
            if isinstance(value, str):
                candidate = value.strip()
                return float(candidate) if candidate else 0.0
            return float(value)
        except (ValueError, TypeError):
            return 0.0

    column_totals = []
    for header in table_headers[2:]:
        col_total = 0.0
        for month_map in month_maps:
            month_info = month_map.get(header, {})
            col_total += _to_float(month_info.get('value', 0))
        column_totals.append(col_total)

    grand_total = 0.0
    for row_data in rows:
        grand_total += _to_float(row_data.get('Total', 0))

    totals_row = table.add_row()
    set_row_height(totals_row, Cm(0.5))
    totals_cells = totals_row.cells

    for col_idx, header in enumerate(table_headers):
        cell = totals_cells[col_idx]
        header_key = str(header).strip().lower()
        is_total_col = header_key == 'total' or (total_col_idx is not None and col_idx == total_col_idx)
        set_cell_shading(cell, '#FFF2CC')
        if col_idx == 0: 
            cell.text = "" if (is_split_table or has_total_in_this_table) else str(len(rows) + 1)
            set_cell_font(cell, DEFAULT_FONT_SIZE, False)
            align_cell(cell, 'center')
            set_cell_vertical_alignment(cell, 'center')
        elif col_idx == 1:  
            cell.text = "Total"
            set_cell_font(cell, STAFF_FONT_SIZE, True)
            align_cell(cell, 'left')
            set_cell_vertical_alignment(cell, 'center')
        elif is_total_col:
            cell.text = str(int(grand_total)) if grand_total == int(grand_total) else str(round(grand_total, 2))
            set_cell_font(cell, DEFAULT_FONT_SIZE, True)
            align_cell(cell, 'center')
            set_cell_vertical_alignment(cell, 'center')
        else:
            column_total = column_totals[col_idx - 2]
            cell.text = str(int(column_total)) if column_total == int(column_total) else str(round(column_total, 2))
            set_cell_font(cell, DEFAULT_FONT_SIZE, True)
            align_cell(cell, 'center')
            set_cell_vertical_alignment(cell, 'center')

    num_static_cols = 2
    num_month_cols = len(month_chunk)
    enforce_column_widths(table, num_static_cols, num_month_cols, has_total_in_this_table)

    if total_col_idx is None:
        total_col_idx = next((i for i, h in enumerate(table_headers) if str(h).strip().lower() == 'total'), None)
    if total_col_idx is not None:
        for row in table.rows:
            if total_col_idx < len(row.cells):
                set_cell_shading(row.cells[total_col_idx], '#FFF2CC')

    return table


def find_all_sdt_by_title(doc: Document, title: str):
    sdts = []
    for el in doc.element.body.iter():
        if el.tag.endswith("sdt"):
            for child in el.iter():
                if child.tag.endswith("alias") and child.get(qn("w:val")) == title:
                    sdts.append(el)
                    break
                if child.tag.endswith("tag") and child.get(qn("w:val")) == title:
                    sdts.append(el)
                    break
    return sdts

def set_table_dotted_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    border_props = {
        'w:val': 'dotted',
        'w:sz': '6',  
        'w:space': '0',
        'w:color': '000000'  
    }
    
    for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right', 'w:insideH', 'w:insideV']:
        border = OxmlElement(border_name)
        for prop, value in border_props.items():
            border.set(qn(prop), value)
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def set_project_brief_table_width(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '4000') 
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    tblJc = OxmlElement('w:jc')
    tblJc.set(qn('w:val'), 'left')
    tblPr.append(tblJc)

def set_project_brief_column_widths(table):
    label_width = Cm(5.0) 
    value_width = Cm(5.0) 
    
    try:
        if len(table.columns) >= 2:
            table.columns[0].width = label_width
            table.columns[1].width = value_width
    except:
        pass
    
    for row in table.rows:
        if len(row.cells) >= 2:
            set_cell_width(row.cells[0], label_width)
            set_cell_width(row.cells[1], value_width)

def generate_project_brief_table(doc: Document, tag: str, data: Dict[str, Any]) -> Document:

    items = data.get("items", [])
    
    if not items:
        return doc

    sdts = find_all_sdt_by_title(doc, tag)
    
    if not sdts:
        return doc
    
    for sdt in sdts:
        sdt_content = next((el for el in sdt.iter() if el.tag.endswith("sdtContent")), None)
        if sdt_content is None:
            continue
        
        for child in list(sdt_content):
            sdt_content.remove(child)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        for item in items:
            label = item.get("label", "")
            value = item.get("value", "")
            
            label_text = f"• {label}"
            
            row = table.add_row()
            cells = row.cells
            cells[0].text = label_text
            cells[1].text = str(value)
            
            align_cell(cells[0], 'left')
            align_cell(cells[1], 'left')
            set_cell_font(cells[0], STAFF_FONT_SIZE, name="Century Gothic")
            set_cell_font(cells[1], STAFF_FONT_SIZE, bold=True, name="Century Gothic")
            set_cell_margins(cells[0], bottom=5)
            set_cell_margins(cells[1], bottom=5)

        set_project_brief_table_width(table)
        set_project_brief_column_widths(table)
        set_table_dotted_border(table)
        
        sdt_content.append(table._tbl)
    
    return doc
